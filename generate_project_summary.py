#!/usr/bin/env python3
"""Generate structured Project Summary.docx from project_summary_content.py."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

from project_summary_content import PROJECT_ORDER, PROJECTS, PROJECT_SHORT

OUT = Path("/workspace/Project Summary.docx")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return h


def add_label_paragraph(doc, label: str, body: str):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}\n")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    run_body = p.add_run(body)
    run_body.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(12)
    return p


def add_bullet_section(doc, label: str, items: list[str]):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_after = Pt(4)
    for item in items:
        bp = doc.add_paragraph(item, style="List Bullet")
        bp.paragraph_format.space_after = Pt(4)
        for run in bp.runs:
            run.font.size = Pt(11)
    doc.add_paragraph()


def add_barrier_impact_section(doc, pairs: list[dict]):
    p = doc.add_paragraph()
    run = p.add_run("The Barriers and the Impacts to the Business due to those Barriers")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p.paragraph_format.space_after = Pt(8)

    for i, pair in enumerate(pairs, 1):
        bp = doc.add_paragraph(style="List Bullet")
        b_run = bp.add_run(f"Barrier: {pair['barrier']}")
        b_run.bold = True
        b_run.font.size = Pt(11)
        bp.paragraph_format.space_after = Pt(2)

        ip = doc.add_paragraph(style="List Bullet 2")
        i_run = ip.add_run(f"Business impact: {pair['impact']}")
        i_run.font.size = Pt(11)
        ip.paragraph_format.space_after = Pt(8)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_heading("GTM Performance & Readiness — Project Summary", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    sub = doc.add_paragraph("Amanda Pattenden | August 2026")
    sub.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sub.paragraph_format.space_after = Pt(18)

    intro = doc.add_paragraph(
        "This document summarises each programme in the GTM Performance & Readiness portfolio. "
        "Content is drawn from approved BRDs, the Requirements Register, and programme delivery experience. "
        "Update as the register is completed."
    )
    intro.paragraph_format.space_after = Pt(18)

    for key in PROJECT_ORDER:
        data = PROJECTS[key]
        short = PROJECT_SHORT.get(key, key)
        doc.add_page_break()
        add_heading(doc, short, level=1)

        add_label_paragraph(doc, "Project Description", data["description"])
        add_label_paragraph(doc, "Project Goal", data["goal"])
        add_bullet_section(doc, "The Benefits to the Business with Successful Delivery", data["benefits"])
        add_barrier_impact_section(doc, data["barrier_impacts"])

    return doc


def main():
    doc = build_document()
    doc.save(str(OUT))
    print(f"Saved {OUT} ({len(PROJECT_ORDER)} projects)")


if __name__ == "__main__":
    main()
